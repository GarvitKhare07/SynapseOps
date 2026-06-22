from io import BytesIO

import fitz
from docx import Document

from backend.services.chunking_service import SourceText


class ExtractionService:
    def extract(self, content: bytes, mime_type: str, filename: str) -> list[SourceText]:
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if mime_type == "application/pdf" or extension == "pdf":
            return self._extract_pdf(content)
        if (
            mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or extension == "docx"
        ):
            return self._extract_docx(content)
        if mime_type.startswith("text/") or extension == "txt":
            return [SourceText(text=self._decode_text(content), page=None)]
        raise ValueError("Unsupported file type")

    def _extract_pdf(self, content: bytes) -> list[SourceText]:
        sources: list[SourceText] = []
        with fitz.open(stream=content, filetype="pdf") as document:
            for index, page in enumerate(document, start=1):
                text = page.get_text("text").strip()
                if text:
                    sources.append(SourceText(text=text, page=index))
        return sources

    def _extract_docx(self, content: bytes) -> list[SourceText]:
        document = Document(BytesIO(content))
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return [SourceText(text="\n".join(parts), page=None)]

    def _decode_text(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")


extraction_service = ExtractionService()
